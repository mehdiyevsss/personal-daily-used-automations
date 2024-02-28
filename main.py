from docx import Document
from docx.shared import RGBColor


doc_path='exdoc.docx'
document=Document(doc_path)

for paragraph in document.paragraphs:
    if '-' in paragraph.text:  
        parts = paragraph.text.split('-', 1)
        first_part = parts[0]

        paragraph.clear()
        run = paragraph.add_run(first_part)
        run.font.color.rgb = RGBColor(255, 0, 0)  

        paragraph.add_run('-' + parts[1])


document.save("exidoc.docx")  
